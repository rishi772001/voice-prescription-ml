'''
@Author: rishi
'''

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm , RGBColor , Inches
from docx.shared import Pt
from tkinter import messagebox


def create(name, age, date, medicines):
    if len(name) == 0:
        messagebox.showinfo("Error", "No patient name is found, please retry again!")
        raise Exception('No name found')
    document = Document()
    para = document.add_heading("")
    run = para.add_run("HealthCare Hospitals \n")
    run = para.add_run(
        " Dr.Rishi Raj M.D.(Neurology) \t Dr.Rishi Kumar M.S.(Ortho) \t Dr.Swetha varnaa M.S.(Ophthalmology) \n No.20,Blossom Avenue,Chennai \n Ph:0452-2588 1522 \n")
    font = run.font
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(153, 17, 150)
    font.size = Pt(16)
    font.color.rgb = RGBColor(217, 17, 213)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    para1 = document.add_paragraph()
    run1 = para1.add_run("Name:{}".format(name))
    run1.font.size = Pt(14)

    para2 = document.add_paragraph()
    run2 = para2.add_run("Age:{}".format(age))
    run2.font.size = Pt(14)

    para3 = document.add_paragraph()
    run3 = para3.add_run("Date:{}".format(date))
    run3.font.size = Pt(14)

    if len(medicines) == 0:
        messagebox.showinfo("Alert", "Please include atleast one medicine")
    else:
        table = document.add_table(len(medicines), 5)

        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Medicine'
        heading_cells[1].text = 'Dosage(mg)'
        heading_cells[2].text = 'Frequency'
        heading_cells[3].text = 'Duration'

        for i in range(len(medicines)):
            cells = table.add_row().cells
            cells[0].text = medicines[i]["DRUG"].capitalize()
            cells[1].text = medicines[i]["STRENGTH"].capitalize()
            cells[2].text = medicines[i]["FREQUENCY"].capitalize()
            cells[3].text = medicines[i]["DURATION"].capitalize()
    # for space

    header = document.sections[0].footer
    f1 = header.add_paragraph()

    runf = f1.add_run("\n Signature \n I hereby accept that this prescription was verified")
    runf.font.size = Pt(16)
    runf.font.color.rgb = RGBColor(0, 0, 0)
    f1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    document.save("{}.docx".format(name))
    messagebox.showinfo("Successfully completed",
                        "Success! Please verify the document generated and click upload \n Document Name:{}.docx".format(
                            name))


if __name__ == "__main__":
    create("kumar", "20", "20-10-2020", [{'DRUG': 'paracetamol', 'STRENGTH': '500mg', 'FREQUENCY': 'thrice', 'DURATION': 'for 5 days'}, {'DRUG': 'Amoxicillin', 'STRENGTH': '650mg', 'FREQUENCY': 'night', 'DURATION': 'for 2 days'}])